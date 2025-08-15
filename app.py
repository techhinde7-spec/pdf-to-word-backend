import os
from flask import Flask, request, send_file
from werkzeug.utils import secure_filename
import fitz  # PyMuPDF
import pytesseract
from PIL import Image
from docx import Document

app = Flask(__name__)
UPLOAD_FOLDER = "uploads"
os.makedirs(UPLOAD_FOLDER, exist_ok=True)

@app.route("/")
def home():
    return "PDF to Word Converter Backend Running ✅"

@app.route("/convert", methods=["POST"])
def convert_pdf():
    if "file" not in request.files:
        return "No file uploaded", 400

    file = request.files["file"]
    filename = secure_filename(file.filename)
    filepath = os.path.join(UPLOAD_FOLDER, filename)
    file.save(filepath)

    # Create Word document
    doc = Document()
    pdf = fitz.open(filepath)

    for page in pdf:
        text = page.get_text()

        if text.strip():  # Text-based PDF
            doc.add_paragraph(text)
        else:  # Scanned PDF → OCR
            pix = page.get_pixmap()
            img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
            ocr_text = pytesseract.image_to_string(img)
            doc.add_paragraph(ocr_text)

    output_path = filepath.replace(".pdf", ".docx")
    doc.save(output_path)
    return send_file(output_path, as_attachment=True)

if __name__ == "__main__":
    app.run(host="0.0.0.0", port=10000)

